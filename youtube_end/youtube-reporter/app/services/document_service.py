import os
import tempfile
from typing import Dict, Any, Optional, Tuple
from fastapi import UploadFile, HTTPException
import PyPDF2
from docx import Document
import pandas as pd
from datetime import datetime

class DocumentService:
    SUPPORTED_EXTENSIONS = {
        '.pdf': 'application/pdf',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        '.csv': 'text/csv',
        '.txt': 'text/plain'
    }

    def __init__(self):
        self.temp_dir = tempfile.mkdtemp()

    def validate_file(self, file: UploadFile) -> Tuple[str, str]:
        """파일 유효성 검사 및 확장자 확인"""
        ext = os.path.splitext(file.filename)[1].lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise HTTPException(
                status_code=400,
                detail=f"지원하지 않는 파일 형식입니다. 지원 형식: {', '.join(self.SUPPORTED_EXTENSIONS.keys())}"
            )
        return ext, self.SUPPORTED_EXTENSIONS[ext]

    async def save_upload_file(self, file: UploadFile) -> str:
        """업로드된 파일을 임시 저장"""
        try:
            ext, _ = self.validate_file(file)
            temp_path = os.path.join(self.temp_dir, f"{datetime.now().strftime('%Y%m%d_%H%M%S')}{ext}")
            
            with open(temp_path, "wb") as buffer:
                content = await file.read()
                buffer.write(content)
            
            return temp_path
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"파일 저장 실패: {str(e)}")

    def extract_text_from_pdf(self, file_path: str) -> Dict[str, Any]:
        """PDF 파일에서 텍스트 추출"""
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                
                return {
                    "text": text,
                    "pages": len(reader.pages),
                    "word_count": len(text.split())
                }
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"PDF 텍스트 추출 실패: {str(e)}")

    def extract_text_from_docx(self, file_path: str) -> Dict[str, Any]:
        """DOCX 파일에서 텍스트 추출"""
        try:
            doc = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            return {
                "text": text,
                "paragraphs": len(doc.paragraphs),
                "word_count": len(text.split())
            }
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"DOCX 텍스트 추출 실패: {str(e)}")

    def extract_text_from_excel(self, file_path: str) -> Dict[str, Any]:
        """Excel 파일에서 텍스트 추출"""
        try:
            df = pd.read_excel(file_path)
            text = df.to_string()
            
            return {
                "text": text,
                "rows": len(df),
                "columns": len(df.columns),
                "word_count": len(text.split())
            }
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Excel 텍스트 추출 실패: {str(e)}")

    def extract_text_from_csv(self, file_path: str) -> Dict[str, Any]:
        """CSV 파일에서 텍스트 추출"""
        try:
            df = pd.read_csv(file_path)
            text = df.to_string()
            
            return {
                "text": text,
                "rows": len(df),
                "columns": len(df.columns),
                "word_count": len(text.split())
            }
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"CSV 텍스트 추출 실패: {str(e)}")

    def extract_text_from_txt(self, file_path: str) -> Dict[str, Any]:
        """TXT 파일에서 텍스트 추출"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            return {
                "text": text,
                "lines": len(text.splitlines()),
                "word_count": len(text.split())
            }
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"TXT 텍스트 추출 실패: {str(e)}")

    async def process_document(self, file: UploadFile) -> Dict[str, Any]:
        """문서 처리 및 텍스트 추출"""
        try:
            # 파일 저장
            file_path = await self.save_upload_file(file)
            ext = os.path.splitext(file_path)[1].lower()
            
            # 파일 크기
            file_size = os.path.getsize(file_path)
            
            # 파일 형식별 텍스트 추출
            if ext == '.pdf':
                result = self.extract_text_from_pdf(file_path)
            elif ext == '.docx':
                result = self.extract_text_from_docx(file_path)
            elif ext == '.xlsx':
                result = self.extract_text_from_excel(file_path)
            elif ext == '.csv':
                result = self.extract_text_from_csv(file_path)
            elif ext == '.txt':
                result = self.extract_text_from_txt(file_path)
            else:
                raise HTTPException(status_code=400, detail="지원하지 않는 파일 형식입니다")
            
            # 임시 파일 삭제
            os.remove(file_path)
            
            return {
                "filename": file.filename,
                "content_type": self.SUPPORTED_EXTENSIONS[ext],
                "size": file_size,
                **result
            }
            
        except Exception as e:
            # 임시 파일 정리
            if os.path.exists(file_path):
                os.remove(file_path)
            raise HTTPException(status_code=500, detail=f"문서 처리 실패: {str(e)}")

document_service = DocumentService() 